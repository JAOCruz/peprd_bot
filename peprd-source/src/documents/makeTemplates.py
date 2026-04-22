#!/usr/bin/env python3
"""
Converts real .docx legal files into docxtemplater templates
by replacing real names/data with {placeholders}.
Run once to generate template files.
"""
import os, re, shutil
from docx import Document
from docx.oxml.ns import qn
import copy

GURU_KB = '/home/jay/.openclaw/workspace/guru-kb/BASE DE DATOS'
OUT_DIR = '/home/jay/.openclaw/workspace/projects/whatsapp-bot/guru-source/templates/generated'

def replace_text_in_runs(paragraph, replacements):
    """Replace text in paragraph runs while preserving formatting."""
    full_text = ''.join(r.text for r in paragraph.runs)
    new_text = full_text
    for old, new in replacements.items():
        new_text = new_text.replace(old, new)
    
    if new_text != full_text and paragraph.runs:
        # Put all text in first run, clear the rest
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''

def make_alquiler_vivienda():
    src = f'{GURU_KB}/CONTRATOS CIVILES/CONTRATOS BAJO FIRMA PRIVADA/RENTAS O ALQUILERES/CONTRATO DE ALQUILER DE VIVIENDA.docx'
    dst = f'{OUT_DIR}/contrato_alquiler_vivienda.docx'
    shutil.copy2(src, dst)
    doc = Document(dst)
    
    replacements = {
        'MIRTHA OZUNA SILVESTRE': '{propietario_nombre}',
        'Mirtha Ozuna Silvestre': '{propietario_nombre}',
        '001-0106890-6': '{propietario_cedula}',
        'GUILLERMINA MARIA MUÑOZ TORRES': '{inquilino_nombre}',
        'Guillermina Maria Muñoz Torres': '{inquilino_nombre}',
        '001-1727916-6': '{inquilino_cedula}',
        'APARTA ESTUDIO DE 1 HABITACION EN EL SEGUNDO PISO DEL ULTIMO EDIF. I, CON TODAS SUS CONSISTENCIAS': '{descripcion_inmueble}',
        'Avenida Independencia #1455, sector Mata Hambre, (EN EL INMUEBLE ALQUILADO), SANTO DOMINGO, DISTRITO NACIONAL, CAPITAL DE LA REPUBLICA DOMINICANA': '{direccion_inmueble}',
        'OCHO MIL PESOS DOMINICANOS (RD$ 8,000.00)': '{monto_alquiler_texto}',
        'DOSCIENTOS PESOS DOMINICANOS (RD$ 200.00)': '{monto_mantenimiento_texto}',
        'VEINTICUATRO MIL PESOS DOMINICANOS (RD$ 24,000.00)': '{monto_deposito_texto}',
        'MICHAEL FAMILIA DE OLEO': '{garante_nombre}',
        'Michael Familia De Oleo': '{garante_nombre}',
        '223-0047820-7': '{garante_cedula}',
        'quince (15) días del mes de marzo del año dos mil veinticuatro (2024)': '{fecha_contrato}',
        'DR. RAMON ANIBAL GUZMAN M.': '{notario_nombre}',
        'DR. RAMON ANIBAL GUZMAN MENDEZ': '{notario_nombre}',
        'No. 4745': 'No. {notario_matricula}',
        # Gender roles - keep as variables for flexibility
        'LA PROPIETARIA': '{propietario_role}',
        'EL PROPIETARIO': '{propietario_role}',
        'LA INQUILINA': '{inquilino_role}',
        'EL INQUILINO': '{inquilino_role}',
        'EL APARTA-ESTUDIO': 'EL INMUEBLE',
        'EL APARTAMENTO': 'EL INMUEBLE',
    }
    
    for para in doc.paragraphs:
        replace_text_in_runs(para, replacements)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_runs(para, replacements)
    
    doc.save(dst)
    print(f'✅ contrato_alquiler_vivienda.docx')

def make_venta_vehiculo():
    src = f'{GURU_KB}/CONTRATOS CIVILES/CONTRATOS BAJO FIRMA PRIVADA/BIENES TRASLATIVOS/VEHICULOS/ACTO DE VENTA DE VEHICULO.docx'
    dst = f'{OUT_DIR}/acto_venta_vehiculo.docx'
    shutil.copy2(src, dst)
    doc = Document(dst)
    
    # Read first to find real names
    text = '\n'.join(p.text for p in doc.paragraphs)
    print(f'Vehicle sale fields preview:\n{text[:600]}\n---')
    doc.save(dst)
    print(f'✅ acto_venta_vehiculo.docx (manual review needed)')

def make_poder_autorizacion():
    src = f'{GURU_KB}/CONTRATOS CIVILES/CONTRATOS AUTENTICOS/PODERES/PODER DE AUTORIZACION AMPLIATORIO MULTIUSO.docx'
    dst = f'{OUT_DIR}/poder_autorizacion.docx'
    shutil.copy2(src, dst)
    doc = Document(dst)
    text = '\n'.join(p.text for p in doc.paragraphs)
    print(f'Poder fields preview:\n{text[:600]}\n---')
    doc.save(dst)
    print(f'✅ poder_autorizacion.docx (manual review needed)')

def make_alquiler_comercial():
    src = f'{GURU_KB}/CONTRATOS CIVILES/CONTRATOS BAJO FIRMA PRIVADA/RENTAS O ALQUILERES/CONTRATO DE ALQUILER LOCAL COMERCIAL.docx'
    dst = f'{OUT_DIR}/contrato_alquiler_comercial.docx'
    shutil.copy2(src, dst)
    doc = Document(dst)
    text = '\n'.join(p.text for p in doc.paragraphs)
    print(f'Commercial lease preview:\n{text[:600]}\n---')
    doc.save(dst)
    print(f'✅ contrato_alquiler_comercial.docx')

if __name__ == '__main__':
    os.makedirs(OUT_DIR, exist_ok=True)
    make_alquiler_vivienda()
    make_venta_vehiculo()
    make_poder_autorizacion()
    make_alquiler_comercial()
    print('\nAll templates generated.')
